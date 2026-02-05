from fastapi.responses import FileResponse
from docx import Document
import tempfile
import os

from app.execution.contracts.execution_request import ExecutionRequest
from app.execution.contracts.execution_context import ExecutionContext


def generate_report(
    request: ExecutionRequest,
    ctx: ExecutionContext,
):
    """
    Minimal v1 report generator.
    Produces a DOCX deterministically.
    """

    doc = Document()
    doc.add_heading("VendorLogic Report", level=1)

    doc.add_paragraph(f"Contract ID: {ctx.contract_id}")
    doc.add_paragraph(f"Execution ID: {ctx.execution_id}")

    doc.add_heading("Status", level=2)
    doc.add_paragraph("Report generated successfully.")

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(tmp.name)

    return FileResponse(
        tmp.name,
        filename="vendorlogic-report.docx",
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
